"""
seed_plantilla_cps.py
────────────────────
Genera el archivo .docx de la plantilla "Contrato de Prestación de Servicios
SaaS – Pymes" y lo registra en la base de datos como PlantillaDocumento.

Uso:
    python manage.py shell < seed_plantilla_cps.py
  o bien:
    python manage.py runscript seed_plantilla_cps  (si usas django-extensions)
  o directamente:
    python seed_plantilla_cps.py  (desde la raíz del proyecto Django)
"""

import os
import sys
import django

# ── Bootstrap Django si se ejecuta directamente ──────────────────────────────
if __name__ == '__main__':
    os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'core.settings')
    django.setup()

# ── Dependencias ──────────────────────────────────────────────────────────────
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.core.files.base import ContentFile
from django.conf import settings
from django.contrib.auth import get_user_model

from catalogo.models import Software
from plantillas.models import PlantillaDocumento

# ── Parámetros ────────────────────────────────────────────────────────────────
NOMBRE_PLANTILLA   = 'Contrato de Prestación de Servicios – SaaS Pyme'
TIPO_CONTRATO      = 'RECURRENTE'   # Coincide con TipoContrato.RECURRENTE
VERSION_CODIGO     = 'v1.0'
SOFTWARE_GLOBAL    = True           # True → software=None (aplica a todos los tenants)
FILENAME_DOCX      = 'CPS_SaaS_Pyme_v1_0.docx'

# ── Helpers de formato ────────────────────────────────────────────────────────

def _set_heading(doc, text, level=1):
    """Agrega un párrafo con estilo Heading N."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # azul CLM
    return h


def _add_body(doc, text):
    """Párrafo de cuerpo normal, justificado."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_var_inline(doc, label, var_name):
    """Párrafo con etiqueta + placeholder Jinja2 para docxtpl."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(f'{label}: ').bold = True
    p.add_run('{{' + var_name + '}}')


def _add_signature_table(doc):
    """Tabla de firmas 2 columnas."""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    headers = ['EL PRESTADOR', 'EL CLIENTE']
    fields  = [
        ('Nombre / Razón Social', ['software.nombre',        'cliente.nombre']),
        ('RUT / ID Fiscal',       ['software.slug',          'cliente.identificador']),
        ('Firma',                 ['_____________________',  '_____________________']),
    ]
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, (label, values) in enumerate(fields, start=1):
        for col_idx, val in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            run_label = p.add_run(f'{label}:\n')
            run_label.bold = True
            run_label.font.size = Pt(9)
            run_val = p.add_run(f'{{% if {val} %}}{{{{{val}}}}}{{% else %}}{val}{{% endif %}}')
            run_val.font.size = Pt(10)


# ── Generación del documento ──────────────────────────────────────────────────

def build_docx() -> bytes:
    doc = Document()

    # ── Estilos globales ──────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Encabezado del documento ──────────────────────────────────────────────
    title = doc.add_heading('CONTRATO DE PRESTACIÓN DE SERVICIOS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub = doc.add_paragraph('Software como Servicio (SaaS) – Modalidad Pyme')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()  # espacio

    # ── Metadatos del contrato ────────────────────────────────────────────────
    meta_table = doc.add_table(rows=3, cols=4)
    meta_data = [
        ('Tipo de Contrato', '{{ contrato.tipo_contrato_display }}',
         'Versión Plantilla', '{{ contrato.id }}'),
        ('Fecha de Inicio',  '{{ contrato.fecha_inicio }}',
         'Fecha de Vencimiento', '{{ contrato.fecha_vencimiento }}'),
        ('Generado el',      '{{ fecha_generacion }}',
         'Estado',           '{{ contrato.etapa_display }}'),
    ]
    for row_idx, row_data in enumerate(meta_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = meta_table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            if col_idx % 2 == 0:
                run = p.add_run(cell_text)
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x7C, 0x76, 0x70)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                run = p.add_run(cell_text)
                run.font.size = Pt(9)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PARTES CONTRATANTES
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '1. PARTES CONTRATANTES')

    _add_body(doc,
        'En la ciudad de ____________, a la fecha indicada en el encabezado del presente instrumento, '
        'comparecen las siguientes partes:')

    # Sub-sección Prestador
    _set_heading(doc, '1.1 El Prestador', level=2)
    _add_body(doc,
        '{{ software.nombre }}, empresa titular del software SaaS aquí contratado, '
        'con domicilio en ____________, identificada con RUT/ID {{ software.slug }}, '
        'representada por quien tiene facultades suficientes para este acto '
        '(en adelante "El Prestador").')

    # Sub-sección Cliente
    _set_heading(doc, '1.2 El Cliente', level=2)
    _add_body(doc,
        '{{ cliente.nombre }}, de tipo {{ cliente.tipo_persona }}, '
        'identificada con RUT/RUN {{ cliente.identificador }}, '
        'con domicilio registrado en la plataforma, '
        'correo electrónico {{ cliente.email }}, '
        'teléfono de contacto {{ cliente.telefono }} '
        '(en adelante "El Cliente").')

    # ══════════════════════════════════════════════════════════════════════════
    # 2. OBJETO
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '2. OBJETO DEL CONTRATO')
    _add_body(doc,
        'El presente instrumento tiene por objeto regular las condiciones bajo las cuales '
        'El Prestador otorgará al Cliente acceso a la plataforma de software {{ software.nombre }} '
        '("el Software"), alojada en infraestructura propia del Prestador (SaaS – Software '
        'como Servicio), incluyendo los servicios de soporte técnico descritos en el Anexo A.')

    _add_body(doc,
        'El Software será accesible vía Internet desde cualquier dispositivo compatible. '
        'El Prestador no transfiere la propiedad ni el código fuente del Software; únicamente '
        'otorga una licencia de uso no exclusiva, intransferible e inembargable durante la '
        'vigencia del contrato.')

    # ══════════════════════════════════════════════════════════════════════════
    # 3. VIGENCIA
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '3. VIGENCIA')
    _add_body(doc,
        'Este Contrato entrará en vigor el día {{ contrato.fecha_inicio }} y permanecerá '
        'vigente hasta el {{ contrato.fecha_vencimiento }}, salvo que alguna de las partes '
        'ejerza su derecho de término anticipado conforme a la Cláusula 10.')

    _add_body(doc,
        'Al término del período pactado, el Contrato se renovará automáticamente por períodos '
        'iguales, salvo aviso escrito en contrario enviado con al menos 30 (treinta) días '
        'calendario de anticipación.')

    # ══════════════════════════════════════════════════════════════════════════
    # 4. PRECIO Y FORMA DE PAGO
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '4. PRECIO Y FORMA DE PAGO')
    _add_body(doc,
        'Las partes acuerdan una tarifa de servicio de {{ contrato.monto_formateado }} '
        '({{ contrato.tipo_contrato_display }}), pagadera mediante transferencia bancaria, '
        'tarjeta de crédito/débito u otro medio electrónico habilitado por El Prestador.')

    _add_body(doc,
        'El Prestador emitirá la factura o boleta correspondiente dentro de los 5 (cinco) '
        'días hábiles siguientes a cada período devengado. El Cliente dispondrá de '
        '{{ contrato.dias_gracia_autorizados }} días de gracia a partir de la fecha de '
        'vencimiento de la factura antes de activarse los mecanismos de mora.')

    _add_body(doc,
        'El no pago oportuno faculta al Prestador a suspender el acceso al Software '
        'sin perjuicio de exigir los perjuicios causados y los intereses legales vigentes.')

    # ══════════════════════════════════════════════════════════════════════════
    # 5. NIVEL DE SERVICIO (SLA)
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '5. NIVEL DE SERVICIO (SLA)')
    _add_body(doc,
        'El Prestador garantiza una disponibilidad mensual del Software de al menos '
        '{{ sla.uptime_garantizado }}% medida sobre el total de horas del mes calendario '
        '("SLA {{ sla.nombre }}"). Se excluyen de este cómputo las ventanas de '
        'mantenimiento programado, los eventos de fuerza mayor y las interrupciones '
        'imputables a proveedores de conectividad del Cliente.')

    _add_body(doc,
        'El tiempo máximo de respuesta ante incidencias críticas es de '
        '{{ sla.tiempo_respuesta_horas }} hora(s). '
        '{{ sla.detalles }}')

    _add_body(doc,
        'En caso de incumplimiento del SLA, El Prestador otorgará al Cliente un crédito '
        'proporcional a las horas de indisponibilidad fuera del umbral garantizado, hasta un '
        'máximo del 30% del valor mensual del servicio.')

    # ══════════════════════════════════════════════════════════════════════════
    # 6. PROPIEDAD INTELECTUAL Y DATOS
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '6. PROPIEDAD INTELECTUAL Y DATOS')
    _add_body(doc,
        'El Software, su código fuente, base de datos, interfaces, algoritmos y toda '
        'documentación asociada son propiedad exclusiva de El Prestador y se encuentran '
        'protegidos por la legislación de propiedad intelectual aplicable.')

    _add_body(doc,
        'Los datos ingresados por el Cliente en el Software ("Datos del Cliente") son de '
        'exclusiva propiedad del Cliente. El Prestador los tratará como información '
        'confidencial y únicamente para los fines del servicio contratado, sin cederlos '
        'ni comercializarlos con terceros.')

    _add_body(doc,
        'Al término del Contrato, El Prestador pondrá a disposición del Cliente un '
        'respaldo exportable de sus Datos durante 30 (treinta) días calendario, tras lo '
        'cual procederá a su eliminación segura.')

    # ══════════════════════════════════════════════════════════════════════════
    # 7. CONFIDENCIALIDAD
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '7. CONFIDENCIALIDAD')
    _add_body(doc,
        'Cada parte se obliga a mantener en estricta confidencialidad toda la información '
        'técnica, comercial, financiera y estratégica que reciba de la otra en virtud de '
        'este Contrato, y a no divulgarla a terceros sin consentimiento escrito previo. '
        'Esta obligación se extenderá por 3 (tres) años posteriores al término del Contrato.')

    # ══════════════════════════════════════════════════════════════════════════
    # 8. PROTECCIÓN DE DATOS PERSONALES
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '8. PROTECCIÓN DE DATOS PERSONALES')
    _add_body(doc,
        'El tratamiento de datos personales de los usuarios del Cliente se regirá por la '
        'Política de Privacidad de El Prestador, disponible en su sitio web, y por la '
        'legislación vigente en materia de protección de datos (Ley N° 19.628 o la norma '
        'que la reemplace). El Prestador actúa como Encargado de Datos en los términos '
        'legales aplicables.')

    # ══════════════════════════════════════════════════════════════════════════
    # 9. LIMITACIÓN DE RESPONSABILIDAD
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '9. LIMITACIÓN DE RESPONSABILIDAD')
    _add_body(doc,
        'La responsabilidad total acumulada del Prestador frente al Cliente, por cualquier '
        'causa y bajo cualquier teoría legal, no excederá el monto total efectivamente '
        'pagado durante los 12 (doce) meses inmediatamente anteriores al evento que origina '
        'la reclamación.')

    _add_body(doc,
        'En ningún caso el Prestador será responsable por daños indirectos, incidentales, '
        'lucro cesante, pérdida de datos o daño emergente, incluso si fue advertido de '
        'tal posibilidad. Esta limitación no aplica en casos de dolo o culpa grave.')

    # ══════════════════════════════════════════════════════════════════════════
    # 10. TÉRMINO ANTICIPADO
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '10. TÉRMINO ANTICIPADO')
    _add_body(doc,
        'Cualquiera de las partes podrá poner término a este Contrato antes de su '
        'vencimiento en los siguientes casos:')

    for item in [
        'a) Incumplimiento material de la otra parte no subsanado dentro de los 15 '
        '   (quince) días hábiles siguientes a la notificación escrita del incumplimiento.',
        'b) Insolvencia, quiebra o cesión de bienes de la otra parte.',
        'c) Por voluntad unilateral, con aviso escrito de 30 (treinta) días calendario '
        '   de anticipación. En este caso, el Cliente no tendrá derecho a reembolso '
        '   del período en curso.',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.space_after = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # 11. TÉRMINOS LEGALES VIGENTES
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '11. TÉRMINOS Y CONDICIONES VIGENTES')
    _add_body(doc,
        'Se incorporan al presente Contrato, como Anexo B, los Términos y Condiciones '
        'Generales de Uso publicados por el Prestador en su versión vigente a la fecha '
        'de suscripción. En caso de contradicción, prevalecerán las cláusulas del '
        'presente instrumento.')

    doc.add_paragraph(
        '{% if terminos_legales.principal %}'
        'Versión de T&C incorporada: {{ terminos_legales.principal.version_codigo }} '
        '(publicada el {{ terminos_legales.principal.fecha_publicacion }}).'
        '{% endif %}'
    )

    # ══════════════════════════════════════════════════════════════════════════
    # 12. RESOLUCIÓN DE DISPUTAS
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, '12. RESOLUCIÓN DE DISPUTAS Y LEY APLICABLE')
    _add_body(doc,
        'Las partes se comprometen a intentar resolver cualquier controversia de manera '
        'amistosa en un plazo de 30 (treinta) días. De no alcanzarse acuerdo, la disputa '
        'se someterá a mediación y, en su defecto, a arbitraje de derecho ante un árbitro '
        'designado por el Centro de Arbitraje y Mediación correspondiente.')

    _add_body(doc,
        'Este Contrato se rige por las leyes vigentes de la República de Chile. '
        'Para todos los efectos legales, las partes fijan su domicilio en la ciudad '
        'de Santiago.')

    # ══════════════════════════════════════════════════════════════════════════
    # FIRMAS
    # ══════════════════════════════════════════════════════════════════════════
    _set_heading(doc, 'FIRMAS')
    _add_body(doc, 'En señal de conformidad con todo lo anteriormente expuesto, las partes '
                   'suscriben el presente Contrato en dos ejemplares del mismo tenor y valor:')
    doc.add_paragraph()
    _add_signature_table(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # ANEXO A – Descripción del Software
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _set_heading(doc, 'ANEXO A – DESCRIPCIÓN DEL SOFTWARE Y SOPORTE')
    _add_body(doc, '{{ software.nombre }}: {{ software.descripcion }}')
    _add_body(doc,
        'Canal de soporte técnico: correo electrónico / portal de tickets habilitado '
        'por El Prestador. Horario de atención estándar: lunes a viernes, 09:00–18:00 h '
        '(hora local Chile Continental), excluyendo feriados legales.')

    # ── Serializar a bytes ────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Registro en base de datos ─────────────────────────────────────────────────

def seed():
    User = get_user_model()

    # Usar el primer superusuario disponible (o None si no hay)
    admin = User.objects.filter(is_superuser=True).first()

    # Verificar si ya existe
    existe = PlantillaDocumento.objects.filter(
        nombre=NOMBRE_PLANTILLA,
        tipo_contrato=TIPO_CONTRATO,
        version_codigo=VERSION_CODIGO,
        software=None,
    ).first()

    if existe:
        print(f'[SKIP] Ya existe la plantilla: "{NOMBRE_PLANTILLA}" ({VERSION_CODIGO})  id={existe.pk}')
        return existe

    docx_bytes = build_docx()

    plantilla = PlantillaDocumento(
        nombre=NOMBRE_PLANTILLA,
        tipo_contrato=TIPO_CONTRATO,
        software=None,       # Global → aplica a todos los tenants/softwares
        version_codigo=VERSION_CODIGO,
        activa=True,
        subida_por=admin,
    )
    plantilla.archivo_docx.save(FILENAME_DOCX, ContentFile(docx_bytes), save=False)
    plantilla.save()

    print(f'[OK] Plantilla creada: "{NOMBRE_PLANTILLA}" ({VERSION_CODIGO})  id={plantilla.pk}')
    print(f'     Archivo: {plantilla.archivo_docx.name}')
    return plantilla


if __name__ == '__main__':
    seed()
