"""Motor de renderizado: plantilla .docx (docxtpl) + contexto -> DocumentoGenerado
(.docx interno + .pdf inmutable descargable).
"""
import hashlib
import io
import subprocess
import tempfile
import uuid
from pathlib import Path

from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from docxtpl import DocxTemplate
from jinja2 import StrictUndefined
from jinja2.sandbox import SandboxedEnvironment
from jinja2.exceptions import TemplateError, UndefinedError, SecurityError

class TolerantUndefined(StrictUndefined):
    """
    Permite el uso de variables como '____________' (líneas de firma) 
    sin arrojar error, asumiendo que son espacios para completar a mano.
    """
    def __str__(self):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return self._undefined_name
        return super().__str__()

    def __html__(self):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return self._undefined_name
        return super().__html__()

    def __bool__(self):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return False
        return super().__bool__()

    def __getattr__(self, name):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return self
        return super().__getattr__(name)

    def __getitem__(self, key):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return self
        return super().__getitem__(key)

    def __call__(self, *args, **kwargs):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return self
        return super().__call__(*args, **kwargs)

    def __iter__(self):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return iter([])
        return super().__iter__()

    def __len__(self):
        if self._undefined_name and self._undefined_name.startswith('___'):
            return 0
        return super().__len__()

from .contexto import construir_contexto
from ..models import DocumentoGenerado, PlantillaDocumento, ModoOrigenPlantilla, Clausula


class PlantillaRenderError(Exception):
    """Error genérico al renderizar una plantilla o convertirla a PDF."""


class VariablesFaltantesError(PlantillaRenderError):
    """La plantilla referencia una variable que no existe en el contexto del contrato."""


class ConversionPDFError(PlantillaRenderError):
    """Falló la conversión docx -> PDF vía LibreOffice."""


class SinPlantillaActivaError(PlantillaRenderError):
    """No hay ninguna PlantillaDocumento activa (ni específica ni global) para el tipo de contrato."""


def construir_docx_desde_clausulas(plantilla=None, contrato=None) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    title = doc.add_heading('CONTRATO DE SERVICIOS (GENERADO POR CLÁUSULAS)', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    
    doc.add_paragraph()

    # ---- TEXTO DE BIENVENIDA / PREÁMBULO ----
    if contrato:
        cliente_nombre = str(contrato.cliente)
        identificador = ""
        persona_juridica = getattr(contrato.cliente, 'personajuridica', None)
        persona_natural = getattr(contrato.cliente, 'personanatural', None)
        if persona_juridica:
            cliente_nombre = persona_juridica.razon_social
            identificador = f"RUT {persona_juridica.rut}"
        elif persona_natural:
            cliente_nombre = persona_natural.nombre_completo
            identificador = f"RUT/RUN {persona_natural.run}"
            
        fecha_str = contrato.fecha_creacion.strftime('%d/%m/%Y') if contrato.fecha_creacion else '_______'
        software_str = contrato.software.nombre if contrato.software else 'el servicio'
        
        intro_text = (
            f"El presente documento (en adelante, el \"Contrato\") se celebra con fecha {fecha_str}, "
            f"entre EL PROVEEDOR, y por la otra parte, {cliente_nombre}{', ' + identificador if identificador else ''} "
            f"(en adelante, el \"Cliente\").\n\n"
            f"Ambas partes reconocen contar con la capacidad legal suficiente para obligarse y acuerdan los siguientes términos "
            f"para la provisión y uso del producto de software {software_str}:"
        )
    else:
        intro_text = (
            "El presente documento (en adelante, el \"Contrato\") se celebra con fecha ____________, "
            "entre EL PROVEEDOR, y por la otra parte, ____________ (en adelante, el \"Cliente\").\n\n"
            "Ambas partes reconocen contar con la capacidad legal suficiente para obligarse y acuerdan los siguientes términos "
            "para la provisión y uso del producto de software ____________:"
        )

    for p_text in intro_text.split('\n'):
        if p_text.strip():
            p_intro = doc.add_paragraph(p_text.strip())
            p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_intro.paragraph_format.space_after = Pt(12)
            
    doc.add_paragraph()
    # ----------------------------------------

    def _numerar_bloques_doc(bloques):
        """Numeración jerárquica para bloques."""
        contadores = [0, 0, 0]
        numeros = []
        for b in bloques:
            nivel = max(0, min(2, int(b.get('nivel') or 0)))
            contadores[nivel] += 1
            for l in range(nivel + 1, 3):
                contadores[l] = 0
            if nivel == 0:
                numero = f"{contadores[0]}."
            elif nivel == 1:
                numero = f"{contadores[0]}.{contadores[1]}"
            else:
                numero = f"{chr(ord('a') + contadores[2] - 1)})"
            numeros.append((numero, nivel))
        return numeros

    def _render_contenido_rico(doc, contenido, indent_pt):
        """Renderiza nodos Tiptap/ProseMirror a docx."""
        if not isinstance(contenido, (dict, list)):
            return
        nodos = contenido.get('content', []) if isinstance(contenido, dict) else contenido
        for nodo in nodos:
            if not isinstance(nodo, dict):
                continue
            tipo = nodo.get('type')
            if tipo == 'paragraph':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = indent_pt
                p.paragraph_format.space_after = Pt(6)
                _agregar_runs(p, nodo.get('content', []))
            elif tipo in ('bulletList', 'orderedList'):
                style = 'List Bullet' if tipo == 'bulletList' else 'List Number'
                for item in nodo.get('content', []):
                    if isinstance(item, dict) and item.get('type') == 'listItem':
                        for hijo in item.get('content', []):
                            if isinstance(hijo, dict) and hijo.get('type') == 'paragraph':
                                p = doc.add_paragraph(style=style)
                                p.paragraph_format.left_indent = indent_pt + Pt(18)
                                p.paragraph_format.space_after = Pt(3)
                                _agregar_runs(p, hijo.get('content', []))

    def _agregar_runs(paragraph, nodos_texto):
        """Agrega runs con formato a un párrafo."""
        for n in nodos_texto:
            if not isinstance(n, dict):
                continue
            if n.get('type') == 'hardBreak':
                paragraph.add_run().add_break(WD_BREAK.LINE)
                continue
            if n.get('type') == 'text':
                run = paragraph.add_run(n.get('text', ''))
                marks = {m.get('type') for m in (n.get('marks') or []) if isinstance(m, dict)}
                if 'bold' in marks:
                    run.bold = True
                if 'italic' in marks:
                    run.italic = True
                if 'underline' in marks:
                    run.underline = True

    def _agregar_bloque(numero, nivel, titulo, texto, contenido):
        """Renderiza un bloque con título, texto plano legado o contenido rico."""
        indent = Pt(18 * nivel)
        if titulo:
            if nivel == 0:
                h = doc.add_heading(f"{numero} {titulo.upper()}", level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
            elif nivel == 1:
                h = doc.add_heading(f"{numero} {titulo.upper()}", level=2)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if h.runs:
                    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                h.paragraph_format.left_indent = indent
            else:  # nivel 2: párrafo en negrita
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = indent
                run = p.add_run(f"{numero} {titulo}")
                run.bold = True
        if contenido and isinstance(contenido, (dict, list)):
            _render_contenido_rico(doc, contenido, indent)
        else:
            for p_text in (texto or '').split('\n'):
                if p_text.strip():
                    p = doc.add_paragraph(p_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.left_indent = indent
                    p.paragraph_format.space_after = Pt(6)

    if contrato and contrato.clausulas_estructuradas:
        numeros_y_niveles = _numerar_bloques_doc(contrato.clausulas_estructuradas)
        for (numero, nivel), bloque in zip(numeros_y_niveles, contrato.clausulas_estructuradas):
            _agregar_bloque(numero, nivel, (bloque.get('titulo') or '').strip(), bloque.get('texto') or '', bloque.get('contenido'))
    elif contrato and contrato.texto_adicional_clausulas:
        for p_text in contrato.texto_adicional_clausulas.split('\n'):
            if p_text.strip():
                # Detectar encabezado si empieza con un número seguido de punto y espacio
                if p_text.strip()[0].isdigit() and '. ' in p_text[:5]:
                    h = doc.add_heading(p_text.strip(), level=1)
                    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if h.runs:
                        h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                else:
                    p = doc.add_paragraph(p_text.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
    else:
        if plantilla:
            clausulas = plantilla.clausulas_seleccionadas.filter(activa=True).prefetch_related('versiones')
        else:
            clausulas = Clausula.objects.filter(activa=True).prefetch_related('versiones')
        
        for i, c in enumerate(clausulas, start=1):
            version = c.versiones.filter(activa=True, tipo='Estándar').first()
            if not version:
                continue
                
            h = doc.add_heading(f"{i}. {c.nombre.upper()}", level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
                
            p = doc.add_paragraph(version.texto)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)

    # ---- TEXTO DE CONCLUSIÓN Y FIRMAS ----
    doc.add_paragraph()
    concl_text = (
        "En señal de conformidad y aceptación de las cláusulas y condiciones estipuladas en el presente Contrato, "
        "las partes lo firman en dos ejemplares de igual tenor y valor, quedando uno en poder de cada parte."
    )
    p_concl = doc.add_paragraph(concl_text)
    p_concl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_concl.paragraph_format.space_before = Pt(12)
    p_concl.paragraph_format.space_after = Pt(36)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p1 = table.rows[0].cells[0].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("________________________________\n")
    p1.add_run("Por EL PROVEEDOR\n\n\n")

    p2 = table.rows[0].cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("________________________________\n")
    if contrato:
        cliente_nombre = str(contrato.cliente)
        if hasattr(contrato.cliente, 'personajuridica'):
            cliente_nombre = contrato.cliente.personajuridica.razon_social
        elif hasattr(contrato.cliente, 'personanatural'):
            cliente_nombre = contrato.cliente.personanatural.nombre_completo
        p2.add_run(f"Por EL CLIENTE\n{cliente_nombre}\n")
    else:
        p2.add_run("Por EL CLIENTE\n____________\n")
    # ----------------------------------------

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def renderizar_docx(plantilla: PlantillaDocumento, contexto: dict, contrato=None) -> bytes:
    """Renderiza la plantilla .docx con el contexto dado. Devuelve los bytes del docx resultante."""
    try:
        if plantilla.modo_origen == ModoOrigenPlantilla.CLAUSULAS:
            docx_base = construir_docx_desde_clausulas(plantilla, contrato)
            doc = DocxTemplate(docx_base)
        else:
            doc = DocxTemplate(plantilla.archivo_docx.path)
        # TolerantUndefined: permite '____________' pero falla si es otra variable no resuelta.
        # SandboxedEnvironment: el texto de cláusulas y de plantillas .docx subidas no es
        # confiable (lo escriben usuarios del tenant) y docxtpl lo evalúa como Jinja2 -
        # el entorno normal permite RCE vía atributos como __class__/__globals__.
        jinja_env = SandboxedEnvironment(undefined=TolerantUndefined)
        doc.render(contexto, jinja_env=jinja_env)
    except UndefinedError as exc:
        raise VariablesFaltantesError(
            f"La plantilla usa una variable que no existe en los datos del contrato: {exc}"
        ) from exc
    except SecurityError as exc:
        raise PlantillaRenderError(f"La plantilla usa una construcción no permitida: {exc}") from exc
    except TemplateError as exc:
        raise PlantillaRenderError(f"Error de sintaxis en la plantilla: {exc}") from exc

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def renderizar_html(plantilla: PlantillaDocumento, contexto: dict) -> bytes:
    """Renderiza la plantilla HTML con el contexto dado. Carga el archivo desde
    DOCS_TEMPLATE_DIR (o templates/), adapta el formato dc de Claude Design a
    página imprimible y evalúa las variables como template de Django."""
    from django.template import engines
    from .html_doc import cargar_plantilla_html, fecha_larga_es
    try:
        contexto = dict(contexto)
        contexto.setdefault('fecha_documento', fecha_larga_es())
        html_imprimible = cargar_plantilla_html(plantilla.ruta_plantilla_html)
        
        clausulas_a_remover = [k for k, v in contexto.items() if v == '__REMOVE__']
        if clausulas_a_remover:
            from bs4 import BeautifulSoup
            import re
            soup = BeautifulSoup(html_imprimible, 'html.parser')
            for c_rem in clausulas_a_remover:
                for node in soup.find_all(string=re.compile(r'\{\{\s*' + re.escape(c_rem) + r'\b')):
                    if node.parent:
                        node.parent.decompose()
                num = c_rem.split('_')[-1]
                t_rem = f'titulo_clausula_{num}'
                for node in soup.find_all(string=re.compile(r'\{\{\s*' + re.escape(t_rem) + r'\b')):
                    if node.parent:
                        node.parent.decompose()
            html_imprimible = str(soup)
            
        html_str = engines['django'].from_string(html_imprimible).render(contexto)
        return html_str.encode('utf-8')
    except Exception as exc:
        raise PlantillaRenderError(f"Error al renderizar plantilla HTML '{plantilla.ruta_plantilla_html}': {exc}") from exc


def convertir_con_libreoffice(entrada_bytes: bytes, ext_origen: str, ext_destino: str, filtro: str = None) -> bytes:
    """Convierte bytes usando LibreOffice headless."""
    binario = getattr(settings, 'LIBREOFFICE_BINARY', 'soffice')

    with tempfile.TemporaryDirectory() as tmpdir:
        entrada_path = Path(tmpdir) / f'entrada.{ext_origen}'
        entrada_path.write_bytes(entrada_bytes)

        perfil_dir = Path(tmpdir) / f'perfil-{uuid.uuid4().hex}'

        args = [
            binario, '--headless', '--norestore',
            f'-env:UserInstallation=file://{perfil_dir}',
            '--convert-to', filtro if filtro else ext_destino, 
            '--outdir', tmpdir, str(entrada_path)
        ]

        try:
            resultado = subprocess.run(
                args, check=True, timeout=60, capture_output=True,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise ConversionPDFError(f"Falló la conversión a {ext_destino} (LibreOffice): {exc}") from exc

        salida_path = entrada_path.with_suffix(f'.{ext_destino}')
        if not salida_path.exists():
            detalle = resultado.stderr.decode(errors='replace') if getattr(resultado, 'stderr', None) else ''
            raise ConversionPDFError(f"LibreOffice no generó el {ext_destino} esperado. {detalle}")

        return salida_path.read_bytes()


def convertir_a_pdf(docx_bytes: bytes) -> bytes:
    """Convierte bytes de un .docx a PDF usando LibreOffice headless."""
    return convertir_con_libreoffice(docx_bytes, 'docx', 'pdf')


def resolver_plantilla_activa(tipo_contrato: str, software_id, tenant) -> PlantillaDocumento:
    """Plantilla activa del tenant específica del software; si no hay, fallback
    a la activa global (sin software) del mismo tenant."""
    especifica = PlantillaDocumento.objects.filter(
        tenant=tenant, tipo_contrato=tipo_contrato, software_id=software_id, activa=True,
    ).first()
    if especifica:
        return especifica

    global_ = PlantillaDocumento.objects.filter(
        tenant=tenant, tipo_contrato=tipo_contrato, software__isnull=True, activa=True,
    ).first()
    if global_:
        return global_

    raise SinPlantillaActivaError(
        f"No hay ninguna plantilla activa para tipo_contrato={tipo_contrato} "
        f"(ni específica del software {software_id} ni global)."
    )


def obtener_preview_pdf(plantilla: PlantillaDocumento) -> Path:
    """PDF de la plantilla tal cual (variables sin resolver), para previsualizar
    en el catálogo. Se cachea en disco por hash del .docx: LibreOffice tarda
    ~1-2 s por conversión y la plantilla no cambia entre requests."""
    from ..models import ModoOrigenPlantilla
    
    es_html = False
    if plantilla.modo_origen == ModoOrigenPlantilla.CLAUSULAS:
        entrada_bytes = construir_docx_desde_clausulas(plantilla).read()
    elif plantilla.modo_origen == ModoOrigenPlantilla.HTML:
        try:
            entrada_bytes = renderizar_html(plantilla, {})
            es_html = True
        except PlantillaRenderError as exc:
            raise ConversionPDFError(f"Error al obtener preview de la plantilla HTML '{plantilla.ruta_plantilla_html}': {exc}")
    else:
        with plantilla.archivo_docx.open('rb') as f:
            entrada_bytes = f.read()
            
    digest = hashlib.sha256(entrada_bytes).hexdigest()[:16]

    cache_dir = Path(settings.MEDIA_ROOT) / 'plantillas_previews'
    cache_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = cache_dir / f'plantilla_{plantilla.id}_{digest}.pdf'

    if not pdf_path.exists():
        if es_html:
            # WeasyPrint y no LibreOffice: respeta flexbox/grid/@page del diseño.
            from .html_doc import html_a_pdf
            pdf_bytes = html_a_pdf(entrada_bytes.decode('utf-8'))
        else:
            pdf_bytes = convertir_a_pdf(entrada_bytes)
        pdf_path.write_bytes(pdf_bytes)
    return pdf_path


def obtener_preview_imagen(plantilla: PlantillaDocumento) -> Path:
    """Extrae la primera página del PDF de la plantilla como imagen (JPG)."""
    import pdfplumber
    pdf_path = obtener_preview_pdf(plantilla)
    img_path = pdf_path.with_suffix('.jpg')
    
    if not img_path.exists():
        with pdfplumber.open(pdf_path) as pdf:
            if pdf.pages:
                page = pdf.pages[0]
                im = page.to_image(resolution=300) # Alta calidad para evitar pixelado
                # Convertir a RGB usando .original que es el objeto de PIL Image
                pil_img = getattr(im, 'original', im)
                if hasattr(pil_img, 'convert'):
                    pil_img = pil_img.convert('RGB')
                pil_img.save(str(img_path), format='JPEG')
    return img_path


def aplicar_campos_manuales(contexto: dict, campos):
    """Mezcla los campos manuales del usuario al contexto de render, ignorando
    variables reservadas del sistema. Un cuerpo_clausula* vacío se marca
    '__REMOVE__' (la sección se elimina del documento). Compartido por la
    generación real y la vista previa de borrador."""
    if not campos:
        return
    from .html_doc import VARIABLES_RESERVADAS
    for clave, valor in campos.items():
        if clave in VARIABLES_RESERVADAS:
            continue
        if str(valor).strip():
            contexto[clave] = str(valor)
        elif clave.startswith('cuerpo_clausula'):
            contexto[clave] = '__REMOVE__'


def generar_documento(contrato, plantilla: PlantillaDocumento = None, usuario=None,
                      campos: dict = None) -> DocumentoGenerado:
    """Orquesta: resuelve plantilla si no viene dada, construye contexto,
    renderiza docx, convierte a PDF, calcula hash y crea el DocumentoGenerado
    (write-once — nunca actualiza un registro existente).

    `campos`: valores manuales para las variables propias de la plantilla HTML
    (ej. para/de/asunto de un memorándum). No pueden pisar los namespaces del
    contrato; quedan auditados en contexto_usado."""
    if plantilla is None:
        plantilla = resolver_plantilla_activa(contrato.tipo_contrato, contrato.software_id, contrato.tenant)

    contexto = construir_contexto(contrato)
    aplicar_campos_manuales(contexto, campos)

    if plantilla.modo_origen == ModoOrigenPlantilla.HTML:
        from .html_doc import html_a_pdf, siguiente_referencia
        # Correlativo de "Referencia" asignado por el sistema (nunca por el
        # usuario ni por `campos`): se consume solo aquí, en generación real
        # -- el preview de la plantilla (obtener_preview_pdf) no pasa por acá,
        # así que no gasta números de la secuencia.
        contexto['referencia'] = siguiente_referencia(contrato.tenant, plantilla.codigo_prefijo)
        html_bytes = renderizar_html(plantilla, contexto)
        # PDF con WeasyPrint (diseño fiel: flexbox/grid/@page). El .docx interno
        # sigue saliendo de LibreOffice, que es el único que escribe Word.
        pdf_bytes = html_a_pdf(html_bytes.decode('utf-8'))
        docx_bytes = convertir_con_libreoffice(html_bytes, 'html', 'docx', 'docx:MS Word 2007 XML')
    else:
        docx_bytes = renderizar_docx(plantilla, contexto, contrato)
        pdf_bytes = convertir_a_pdf(docx_bytes)
        
    hash_pdf = hashlib.sha256(pdf_bytes).hexdigest()

    # El contexto se serializa para auditoría: valores no-JSON-nativos (Decimal, date)
    # se convierten a texto para poder guardarlos en el JSONField.
    contexto_serializable = _serializar_contexto(contexto)

    with transaction.atomic():
        documento = DocumentoGenerado.objects.create(
            contrato=contrato,
            plantilla=plantilla,
            hash_sha256=hash_pdf,
            contexto_usado=contexto_serializable,
            generado_por=usuario if (usuario and usuario.is_authenticated) else None,
        )
        nombre_base = f"contrato_{contrato.id}_{plantilla.version_codigo}_{documento.id}"
        documento.archivo_docx.save(f"{nombre_base}.docx", ContentFile(docx_bytes), save=False)
        documento.archivo_pdf.save(f"{nombre_base}.pdf", ContentFile(pdf_bytes), save=False)
        documento.save(update_fields=['archivo_docx', 'archivo_pdf'])

    return documento


def _serializar_contexto(contexto: dict):
    """Convierte recursivamente Decimal/date/etc. a tipos JSON-nativos para
    poder guardar el contexto usado como snapshot de auditoría."""
    import decimal
    import datetime

    if isinstance(contexto, dict):
        return {k: _serializar_contexto(v) for k, v in contexto.items()}
    if isinstance(contexto, list):
        return [_serializar_contexto(v) for v in contexto]
    if isinstance(contexto, (decimal.Decimal, datetime.date, datetime.datetime)):
        return str(contexto)
    return contexto
