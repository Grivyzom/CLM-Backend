"""
Servicios de exportación: Django models → archivos (Excel, Word, PDF).
Cada función recibe datos y devuelve BytesIO listo para HttpResponse.
"""
import csv
import io
import zipfile
from datetime import date
from xml.sax.saxutils import escape as xml_escape

import openpyxl
from openpyxl.packaging.custom import StringProperty
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, Cm
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer


# ─── METADATOS DE AUDITORÍA (OpenXML docProps) ───────────────────────────────

def _aplicar_metadata_auditoria(wb, meta):
    """
    Puebla las propiedades nativas del .xlsx (docProps/core.xml + custom.xml)
    con datos de trazabilidad: quién, cuándo, desde qué IP y con qué filtros se
    exportó. Visible en Windows/Explorador vía clic derecho → Propiedades → Detalles.
    """
    autor = f"{meta['usuario_nombre']} (ID: {meta['usuario_id']})"
    wb.properties.creator = autor
    wb.properties.lastModifiedBy = meta['usuario_nombre']
    wb.properties.title = meta['titulo']

    # "Comentarios" (dc:description → Windows lo muestra como "Comentarios")
    wb.properties.description = (
        f"Rol del Emisor: {meta['rol']}\n"
        f"IP de la Sesión: {meta['ip']}\n"
        f"Entorno y Versión: {meta['entorno']}\n"
        f"Filtros Aplicados: {meta['filtros']}"
    )

    for nombre, valor in [
        ('RolEmisor', meta['rol']),
        ('IPSesion', meta['ip']),
        ('EntornoVersion', meta['entorno']),
        ('FiltrosAplicados', meta['filtros']),
    ]:
        wb.custom_doc_props.append(StringProperty(name=nombre, value=str(valor)))


def _inyectar_company(buf, company):
    """
    openpyxl no expone 'Company' vía API pública (vive en docProps/app.xml,
    propiedades extendidas que openpyxl escribe siempre vacías). Se reescribe
    ese único entry dentro del .xlsx (que es un zip) después de guardar.
    """
    buf.seek(0)
    zin = zipfile.ZipFile(buf, 'r')
    out = io.BytesIO()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == 'docProps/app.xml' and b'<Company>' not in data:
                texto = data.decode('utf-8').replace(
                    '</Properties>', f'<Company>{xml_escape(company)}</Company></Properties>'
                )
                data = texto.encode('utf-8')
            zout.writestr(item, data)
    out.seek(0)
    return out


# ─── EXCEL ────────────────────────────────────────────────────────────────────

def _estilo_encabezado(ws, fila=1):
    fill = PatternFill("solid", fgColor="1F3864")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[fila]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center")


def contratos_a_excel(queryset, meta=None):
    """
    Exporta QuerySet de Contrato a Excel. Devuelve BytesIO.
    Si se pasa `meta` (ver documentos.services.auditoria.build_audit_meta), se
    inyectan metadatos de auditoría en las propiedades nativas del documento.
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Contratos"

    encabezados = [
        "ID", "Cliente", "Software", "SLA",
        "Tipo", "Estado", "Monto",
        "Fecha Inicio", "Fecha Vencimiento", "Días Gracia",
    ]
    ws.append(encabezados)
    _estilo_encabezado(ws)

    for c in queryset.select_related('cliente', 'software', 'sla'):
        nombre_cliente = str(c.cliente)
        ws.append([
            c.id,
            nombre_cliente,
            c.software.nombre,
            c.sla.nombre,
            c.get_tipo_contrato_display(),
            c.get_status_display(),
            float(c.monto),
            c.fecha_inicio,
            c.fecha_vencimiento,
            c.dias_gracia_autorizados,
        ])

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 4

    if meta:
        _aplicar_metadata_auditoria(wb, meta)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    if meta and meta.get('company'):
        buf = _inyectar_company(buf, meta['company'])

    return buf


def clientes_a_excel(queryset, meta=None):
    """
    Exporta QuerySet de Cliente a Excel. Devuelve BytesIO.
    Si se pasa `meta` (ver documentos.services.auditoria.build_audit_meta), se
    inyectan metadatos de auditoría en las propiedades nativas del documento.
    """
    from clientes.models import PersonaNatural, PersonaJuridica

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Clientes"

    encabezados = [
        "ID", "Tipo", "Identificador", "Nombre / Razón Social",
        "Email", "Teléfono", "Activo", "Fecha Registro",
    ]
    ws.append(encabezados)
    _estilo_encabezado(ws)

    for cliente in queryset:
        tipo = "—"
        identificador = "—"
        nombre = str(cliente)

        if hasattr(cliente, 'personanatural'):
            tipo = "Persona Natural"
            identificador = cliente.personanatural.run
        elif hasattr(cliente, 'personajuridica'):
            tipo = "Persona Jurídica"
            identificador = cliente.personajuridica.rut

        ws.append([
            cliente.id,
            tipo,
            identificador,
            nombre,
            cliente.email_principal,
            cliente.telefono_contacto or "",
            "Sí" if cliente.is_active else "No",
            cliente.fecha_registro.strftime("%Y-%m-%d %H:%M") if cliente.fecha_registro else "",
        ])

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 4

    if meta:
        _aplicar_metadata_auditoria(wb, meta)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    if meta and meta.get('company'):
        buf = _inyectar_company(buf, meta['company'])

    return buf


def clientes_a_csv(queryset):
    """Exporta iterable de Cliente a CSV. Devuelve BytesIO (UTF-8 con BOM para Excel)."""
    encabezados = [
        "ID", "Tipo", "Identificador", "Nombre / Razón Social",
        "Email", "Teléfono", "Activo", "Fecha Registro",
    ]

    text_buf = io.StringIO()
    writer = csv.writer(text_buf)
    writer.writerow(encabezados)

    for cliente in queryset:
        tipo = "—"
        identificador = "—"
        nombre = str(cliente)

        if hasattr(cliente, 'personanatural'):
            tipo = "Persona Natural"
            identificador = cliente.personanatural.run
        elif hasattr(cliente, 'personajuridica'):
            tipo = "Persona Jurídica"
            identificador = cliente.personajuridica.rut

        writer.writerow([
            cliente.id,
            tipo,
            identificador,
            nombre,
            cliente.email_principal,
            cliente.telefono_contacto or "",
            "Sí" if cliente.is_active else "No",
            cliente.fecha_registro.strftime("%Y-%m-%d %H:%M") if cliente.fecha_registro else "",
        ])

    buf = io.BytesIO(b"\xef\xbb\xbf" + text_buf.getvalue().encode("utf-8"))
    buf.seek(0)
    return buf


# ─── WORD ─────────────────────────────────────────────────────────────────────

def contrato_a_word(contrato):
    """Genera documento Word de un Contrato. Devuelve BytesIO."""
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Título
    titulo = doc.add_heading("CONTRATO DE SERVICIOS", 0)
    titulo.alignment = 1  # center

    doc.add_paragraph(f"Contrato N.° {contrato.id}").runs[0].bold = True
    doc.add_paragraph(f"Fecha de emisión: {date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph()

    # Partes
    doc.add_heading("1. PARTES", level=1)
    doc.add_paragraph(f"Proveedor del Servicio: {contrato.software.nombre}")
    doc.add_paragraph(f"Cliente: {contrato.cliente}")
    doc.add_paragraph(f"Contacto: {contrato.cliente.email_principal}")
    doc.add_paragraph()

    # Condiciones
    doc.add_heading("2. CONDICIONES DEL CONTRATO", level=1)
    tabla_datos = [
        ("Tipo de Contrato", contrato.get_tipo_contrato_display()),
        ("Estado", contrato.get_status_display()),
        ("Monto", f"${float(contrato.monto):,.4f}"),
        ("Fecha de Inicio", contrato.fecha_inicio.strftime("%d/%m/%Y") if contrato.fecha_inicio else "—"),
        ("Fecha de Vencimiento", contrato.fecha_vencimiento.strftime("%d/%m/%Y") if contrato.fecha_vencimiento else "Indefinido"),
        ("Días de Gracia", str(contrato.dias_gracia_autorizados)),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Valor"
    for campo, valor in tabla_datos:
        row = table.add_row().cells
        row[0].text = campo
        row[1].text = valor

    doc.add_paragraph()

    # SLA
    doc.add_heading("3. NIVEL DE SERVICIO (SLA)", level=1)
    sla = contrato.sla
    doc.add_paragraph(f"Nombre SLA: {sla.nombre}")
    doc.add_paragraph(f"Uptime garantizado: {sla.uptime_garantizado}%")
    doc.add_paragraph(f"Tiempo de respuesta: {sla.tiempo_respuesta_horas} horas")
    if sla.detalles:
        doc.add_paragraph(f"Detalles: {sla.detalles}")

    doc.add_paragraph()
    doc.add_paragraph("___________________________          ___________________________")
    doc.add_paragraph("Firma Proveedor                           Firma Cliente")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF ──────────────────────────────────────────────────────────────────────

def contrato_a_pdf(contrato):
    """Genera PDF de un Contrato. Devuelve BytesIO."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=3*cm, rightMargin=2.5*cm,
                            topMargin=2.5*cm, bottomMargin=2.5*cm)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        "titulo", parent=styles["Title"],
        fontSize=16, spaceAfter=6,
    )
    story.append(Paragraph("CONTRATO DE SERVICIOS", title_style))
    story.append(Paragraph(f"Contrato N.° {contrato.id}", styles["Normal"]))
    story.append(Paragraph(f"Fecha de emisión: {date.today().strftime('%d/%m/%Y')}", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("1. PARTES", styles["Heading2"]))
    story.append(Paragraph(f"Proveedor: {contrato.software.nombre}", styles["Normal"]))
    story.append(Paragraph(f"Cliente: {contrato.cliente}", styles["Normal"]))
    story.append(Paragraph(f"Email: {contrato.cliente.email_principal}", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("2. CONDICIONES DEL CONTRATO", styles["Heading2"]))
    data = [
        ["Campo", "Valor"],
        ["Tipo", contrato.get_tipo_contrato_display()],
        ["Estado", contrato.get_status_display()],
        ["Monto", f"${float(contrato.monto):,.4f}"],
        ["Inicio", contrato.fecha_inicio.strftime("%d/%m/%Y") if contrato.fecha_inicio else "—"],
        ["Vencimiento", contrato.fecha_vencimiento.strftime("%d/%m/%Y") if contrato.fecha_vencimiento else "Indefinido"],
        ["Días Gracia", str(contrato.dias_gracia_autorizados)],
    ]
    t = Table(data, colWidths=[6*cm, 10*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF0F8")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    story.append(Paragraph("3. NIVEL DE SERVICIO (SLA)", styles["Heading2"]))
    sla = contrato.sla
    sla_data = [
        ["SLA", sla.nombre],
        ["Uptime garantizado", f"{sla.uptime_garantizado}%"],
        ["Tiempo de respuesta", f"{sla.tiempo_respuesta_horas} horas"],
    ]
    if sla.detalles:
        sla_data.append(["Detalles", sla.detalles])
    t2 = Table(sla_data, colWidths=[6*cm, 10*cm])
    t2.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#EBF0F8")]),
        ("PADDING", (0, 0), (-1, -1), 6),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
    ]))
    story.append(t2)

    doc.build(story)
    buf.seek(0)
    return buf


def reporte_contratos_pdf(queryset, meta=None, pagina_info=None):
    """
    Genera PDF con listado de contratos (una página de resultados, ver
    `pagina_info`). Devuelve BytesIO.

    `queryset` debe venir YA recortado (slice) por el caller — esta función no
    pagina, solo renderiza lo que recibe. Cargar 50k filas en un solo Table sin
    columnas de ancho fijo es lo que hacía lento el reporte completo (reportlab
    mide cada celda para autocalcular anchos); por eso acá siempre se fija
    `colWidths` y se espera un recorte razonable de filas (ver
    documentos.views.exportar_reporte_contratos_pdf para la paginación real).

    Si se pasa `meta`, se puebla el diccionario /Info nativo del PDF (Author,
    Title, Subject, Creator, Keywords) con los datos de auditoría — visible en
    cualquier lector de PDF vía Propiedades del documento.

    `pagina_info`: {'page', 'page_size', 'total', 'total_pages'} — se muestra
    como encabezado ("Página X de Y — mostrando N de TOTAL registros").
    """
    buf = io.BytesIO()
    doc_kwargs = {}
    if meta:
        doc_kwargs = dict(
            title=meta['titulo'],
            author=f"{meta['usuario_nombre']} (ID: {meta['usuario_id']})",
            subject=meta['filtros'],
            creator=meta.get('company', 'Enfoque Platform'),
            keywords=[
                f"Rol:{meta['rol']}",
                f"IP:{meta['ip']}",
                f"Entorno:{meta['entorno']}",
            ],
        )
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm,
                            **doc_kwargs)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("REPORTE DE CONTRATOS", styles["Title"]))
    story.append(Paragraph(f"Generado: {date.today().strftime('%d/%m/%Y')}", styles["Normal"]))
    if pagina_info:
        desde = (pagina_info['page'] - 1) * pagina_info['page_size'] + 1
        hasta = min(pagina_info['page'] * pagina_info['page_size'], pagina_info['total'])
        story.append(Paragraph(
            f"Página {pagina_info['page']} de {pagina_info['total_pages']} "
            f"— mostrando {desde}–{hasta} de {pagina_info['total']} registros",
            styles["Normal"],
        ))
    story.append(Spacer(1, 0.5*cm))

    data = [["ID", "Cliente", "Software", "Tipo", "Estado", "Monto", "Vencimiento"]]
    for c in queryset.select_related('cliente', 'software', 'sla'):
        data.append([
            str(c.id),
            str(c.cliente)[:30],
            c.software.nombre[:20],
            c.get_tipo_contrato_display(),
            c.get_status_display(),
            f"${float(c.monto):,.2f}",
            c.fecha_vencimiento.strftime("%d/%m/%Y") if c.fecha_vencimiento else "—",
        ])

    # Anchos fijos: sin esto reportlab mide cada celda de cada fila para
    # autocalcular columnas, lo que vuelve el reporte casi inutilizable
    # apenas la tabla crece a miles de filas.
    col_widths = [1.3*cm, 4.5*cm, 3.2*cm, 2.3*cm, 2.3*cm, 2.2*cm, 2.2*cm]

    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF0F8")]),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)

    doc.build(story)
    buf.seek(0)
    return buf
