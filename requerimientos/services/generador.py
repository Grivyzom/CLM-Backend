"""
Generación de documentos de Toma de Requerimientos: Requerimiento (preguntas +
respuestas) → .docx / .pdf. Generación nativa (python-docx / reportlab), sin
LibreOffice: acá no hay un .docx de origen subido por el usuario que convertir,
las "preguntas" de la plantilla son la fuente de verdad y el documento se
compone directamente desde ellas (mismo estilo que
documentos/services/exportar.py::contrato_a_word/contrato_a_pdf).
"""
import hashlib
import io
from datetime import date
from xml.sax.saxutils import escape as _xml_escape

from django.core.files.base import ContentFile
from django.db import transaction

from docx import Document
from docx.shared import Cm
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from .. import models as m


class GeneracionExistenteError(Exception):
    """Ya existe un documento generado para este Requerimiento; requiere forzar=True."""


def _nombre_cliente(cliente):
    return (
        getattr(getattr(cliente, 'personajuridica', None), 'razon_social', None)
        or getattr(getattr(cliente, 'personanatural', None), 'nombre_completo', None)
        or str(cliente)
    )


def _esc(texto):
    """Escapa texto para insertarlo en un reportlab.platypus.Paragraph.

    Paragraph interpreta su contenido como un mini-XML (soporta <img src="...">,
    que reportlab resuelve con urlopen si no es un archivo local) - cualquier
    dato de usuario (razon_social, texto de pregunta, etc.) debe escaparse antes
    de pasar por acá para evitar SSRF/lectura de archivos locales.
    """
    return _xml_escape(str(texto))


def _valor_mostrable(pregunta, respuestas):
    valor = respuestas.get(pregunta['id'])
    if valor in (None, ''):
        return '—'
    if pregunta.get('tipo') == 'booleano':
        return 'Sí' if valor in (True, 'true', 'True', '1', 1) else 'No'
    return str(valor)


# ─── WORD ─────────────────────────────────────────────────────────────────────

def generar_docx(requerimiento):
    """Genera el .docx de un Requerimiento. Devuelve BytesIO."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    titulo = doc.add_heading("TOMA DE REQUERIMIENTOS", 0)
    titulo.alignment = 1  # center

    doc.add_paragraph(f"Requerimiento N.° {requerimiento.id}").runs[0].bold = True
    doc.add_paragraph(f"Fecha: {date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph()

    doc.add_heading("1. DATOS GENERALES", level=1)
    doc.add_paragraph(f"Cliente: {_nombre_cliente(requerimiento.cliente)}")
    if requerimiento.contrato_id:
        doc.add_paragraph(f"Contrato asociado: CTR-{str(requerimiento.contrato_id).zfill(6)}")
    doc.add_paragraph(f"Categoría de software: {requerimiento.categoria_producto}")
    doc.add_paragraph(f"Plantilla utilizada: {requerimiento.plantilla.nombre}")
    doc.add_paragraph()

    respuestas = requerimiento.respuestas or {}
    for i, seccion in enumerate(requerimiento.plantilla.secciones, start=2):
        doc.add_heading(f"{i}. {seccion.get('titulo', '').upper()}", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Pregunta"
        hdr[1].text = "Respuesta"
        for pregunta in seccion.get('preguntas', []):
            row = table.add_row().cells
            row[0].text = pregunta.get('texto', '')
            row[1].text = _valor_mostrable(pregunta, respuestas)
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── PDF ──────────────────────────────────────────────────────────────────────

def generar_pdf(requerimiento):
    """Genera el PDF de un Requerimiento. Devuelve BytesIO."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=3 * cm, rightMargin=2.5 * cm,
                            topMargin=2.5 * cm, bottomMargin=2.5 * cm)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle("titulo", parent=styles["Title"], fontSize=16, spaceAfter=6)
    story.append(Paragraph("TOMA DE REQUERIMIENTOS", title_style))
    story.append(Paragraph(f"Requerimiento N.° {requerimiento.id}", styles["Normal"]))
    story.append(Paragraph(f"Fecha: {date.today().strftime('%d/%m/%Y')}", styles["Normal"]))
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph("1. DATOS GENERALES", styles["Heading2"]))
    story.append(Paragraph(f"Cliente: {_esc(_nombre_cliente(requerimiento.cliente))}", styles["Normal"]))
    if requerimiento.contrato_id:
        story.append(Paragraph(
            f"Contrato asociado: CTR-{str(requerimiento.contrato_id).zfill(6)}", styles["Normal"]))
    story.append(Paragraph(f"Categoría de software: {_esc(requerimiento.categoria_producto)}", styles["Normal"]))
    story.append(Spacer(1, 0.5 * cm))

    respuestas = requerimiento.respuestas or {}
    for i, seccion in enumerate(requerimiento.plantilla.secciones, start=2):
        story.append(Paragraph(f"{i}. {_esc(seccion.get('titulo', '').upper())}", styles["Heading2"]))
        data = [["Pregunta", "Respuesta"]]
        for pregunta in seccion.get('preguntas', []):
            data.append([pregunta.get('texto', ''), _valor_mostrable(pregunta, respuestas)])
        t = Table(data, colWidths=[7 * cm, 9 * cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF0F8")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5 * cm))

    doc.build(story)
    buf.seek(0)
    return buf


# ─── ORQUESTACIÓN ─────────────────────────────────────────────────────────────

def generar_documento(requerimiento, usuario=None, forzar=False):
    """Genera y persiste el documento (docx+pdf) de un Requerimiento.

    Registro write-once: si ya existe una generación previa y no viene
    forzar=True, se rechaza (mismo guard que
    plantillas/services/renderizado.py::generar_documento) — no se pisa
    silenciosamente un documento ya emitido/entregado."""
    if m.RequerimientoGenerado.objects.filter(requerimiento=requerimiento).exists() and not forzar:
        raise GeneracionExistenteError(
            "Ya existe un documento generado para este requerimiento. "
            "Confirma para generar una nueva versión (no elimina la anterior)."
        )

    docx_buf = generar_docx(requerimiento)
    pdf_buf = generar_pdf(requerimiento)
    docx_bytes = docx_buf.getvalue()
    pdf_bytes = pdf_buf.getvalue()
    hash_pdf = hashlib.sha256(pdf_bytes).hexdigest()

    with transaction.atomic():
        documento = m.RequerimientoGenerado.objects.create(
            requerimiento=requerimiento,
            hash_sha256=hash_pdf,
            generado_por=usuario,
        )
        nombre_base = f"requerimiento_{requerimiento.id}_{documento.id}"
        documento.archivo_docx.save(f"{nombre_base}.docx", ContentFile(docx_bytes), save=False)
        documento.archivo_pdf.save(f"{nombre_base}.pdf", ContentFile(pdf_bytes), save=False)
        documento.save(update_fields=['archivo_docx', 'archivo_pdf'])

        requerimiento.estado = m.EstadoRequerimiento.GENERADO
        requerimiento.fecha_generacion = documento.fecha_generacion
        requerimiento.save(update_fields=['estado', 'fecha_generacion'])

    return documento
